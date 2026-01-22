from . import Document, Page, DocumentLoader
from docx import Document as DocxDocument
from .image_utils import image_to_base64_uri
from .vision_utils import process_image_with_vision

class DocxLoader(DocumentLoader):
    """A loader for Microsoft Word (.docx) files."""

    def load(self, source: str) -> Document:
        """Reads text and extracts images from a .docx file."""
        try:
            word_doc = DocxDocument(source)
            page_elements = []

            # Extract paragraphs
            for para in word_doc.paragraphs:
                if para.text.strip():
                    page_elements.append({
                        "type": "text",
                        "content": para.text
                    })

            # Extract images
            for rel in word_doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_bytes = rel.target_part.blob
                        base64_str = image_to_base64_uri(image_bytes)
                        
                        if base64_str:
                            vision_result = process_image_with_vision(image_bytes)

                            page_elements.append({
                                "type": "image",
                                "base64": base64_str,
                                "vision_description": vision_result["description"],
                                "vision_tokens": vision_result["total_tokens"],
                                "vision_cost": vision_result["cost"]
                            })
                    except Exception as e:
                        print(f"Warning: Could not process an image: {e}")

            single_page = Page(page_number=1, structured_elements=page_elements)
            
            print(f"✅ Successfully loaded DOCX file")
            return Document(source=source, pages=[single_page])
            
        except Exception as e:
            print(f"Error reading DOCX file: {str(e)}")
            raise e
